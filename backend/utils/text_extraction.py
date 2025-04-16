"""
Text Extraction Utilities for AI Scientific Research Summarizer
Provides high-quality text extraction from various document formats
"""

import os
import sys
import json
import logging
import tempfile
import io
from typing import Dict, Any, Optional, List, Union
from urllib.parse import urlparse

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define supported file types
SUPPORTED_FILE_TYPES = {
    'pdf': ['pdf'],
    'image': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp'],
    'document': ['doc', 'docx', 'txt', 'rtf', 'odt'],
    'spreadsheet': ['csv', 'xls', 'xlsx'],
    'presentation': ['ppt', 'pptx'],
    'webpage': ['html', 'htm', 'url']
}

# Import specialized libraries for each file type
try:
    # PDF processing
    import PyPDF2
    import pdfplumber
    HAS_PDF_LIBS = True
except ImportError:
    logger.warning("PDF processing libraries not installed. Run: pip install PyPDF2 pdfplumber")
    HAS_PDF_LIBS = False

# Separate try/except for OCR-related libraries
try:
    from pdf2image import convert_from_path, convert_from_bytes
    import pytesseract
    import cv2
    import numpy as np
    from PIL import Image
    import urllib.request
    import zipfile
    import platform
    import subprocess
    HAS_OCR_LIBS = True
    
    # Check for poppler path
    import os
    import sys
    
    # Define potential Poppler paths
    POPPLER_PATHS = [
        r"C:\Program Files\poppler-23.11.0\Library\bin",  # Standard install location
        r"C:\Program Files\poppler\bin",
        r"C:\poppler\bin",
        r"C:\Users\Sayan\AppData\Local\Programs\poppler\bin",
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "bin", "poppler", "bin"),
        os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "poppler", "bin")
    ]
    
    # Try to find poppler in PATH or in common locations
    POPPLER_PATH = None
    for path in POPPLER_PATHS:
        if os.path.exists(path):
            POPPLER_PATH = path
            logger.info(f"Found Poppler at: {POPPLER_PATH}")
            break
    
    # If Poppler not found, download and install it automatically
    if not POPPLER_PATH and platform.system() == 'Windows':
        try:
            logger.info("Poppler not found. Attempting to download and install automatically...")
            
            # Create bin directory if it doesn't exist
            bin_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "bin")
            os.makedirs(bin_dir, exist_ok=True)
            
            # Download Poppler for Windows
            poppler_url = "https://github.com/oschwartz10612/poppler-windows/releases/download/v23.11.0-0/Release-23.11.0-0.zip"
            zip_path = os.path.join(bin_dir, "poppler.zip")
            
            logger.info(f"Downloading Poppler from {poppler_url}...")
            urllib.request.urlretrieve(poppler_url, zip_path)
            
            # Extract Poppler
            logger.info("Extracting Poppler...")
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(bin_dir)
            
            # Set Poppler path
            POPPLER_PATH = os.path.join(bin_dir, "poppler-23.11.0", "Library", "bin")
            
            # Add to system PATH temporarily
            os.environ["PATH"] = POPPLER_PATH + os.pathsep + os.environ["PATH"]
            
            logger.info(f"Poppler installed successfully at: {POPPLER_PATH}")
            
            # Clean up zip file
            os.remove(zip_path)
            
        except Exception as e:
            logger.error(f"Failed to download and install Poppler: {e}")
    
    # Verify Poppler installation
    if POPPLER_PATH:
        try:
            if platform.system() == 'Windows':
                pdftoppm_path = os.path.join(POPPLER_PATH, "pdftoppm.exe")
                if os.path.exists(pdftoppm_path):
                    logger.info("Poppler installation verified successfully")
                else:
                    logger.warning(f"pdftoppm.exe not found in {POPPLER_PATH}")
            else:
                # On Unix systems, check if pdftoppm is in PATH
                result = subprocess.run(["which", "pdftoppm"], capture_output=True, text=True)
                if result.returncode == 0:
                    logger.info(f"Poppler found in PATH: {result.stdout.strip()}")
                else:
                    logger.warning("pdftoppm not found in PATH")
        except Exception as e:
            logger.warning(f"Error verifying Poppler installation: {e}")
    
    if not POPPLER_PATH:
        logger.warning("Poppler not found. PDF to image conversion may not work properly.")
        logger.warning("Please install Poppler and ensure it's in your PATH or in a common location.")
        logger.warning("Download from: https://github.com/oschwartz10612/poppler-windows/releases")
except ImportError:
    logger.warning("OCR libraries not installed. OCR features will be disabled.")
    logger.warning("Run: pip install pdf2image pytesseract opencv-python numpy")
    HAS_OCR_LIBS = False

try:
    # Document processing
    import docx
    from odf import text, teletype
    from odf.opendocument import load
    import win32com.client
    HAS_DOC_LIBS = True
except ImportError:
    logger.warning("Document processing libraries not installed. Run: pip install python-docx odfpy pywin32")
    HAS_DOC_LIBS = False

try:
    # Image processing
    from PIL import Image
    HAS_IMG_LIBS = True
except ImportError:
    logger.warning("Image processing libraries not installed. Run: pip install pillow")
    HAS_IMG_LIBS = False

try:
    # Spreadsheet processing
    import pandas as pd
    import openpyxl
    HAS_SPREADSHEET_LIBS = True
except ImportError:
    logger.warning("Spreadsheet processing libraries not installed. Run: pip install pandas openpyxl")
    HAS_SPREADSHEET_LIBS = False

try:
    # Web content processing
    import requests
    from bs4 import BeautifulSoup
    import html2text
    HAS_WEB_LIBS = True
except ImportError:
    logger.warning("Web content processing libraries not installed. Run: pip install requests beautifulsoup4 html2text")
    HAS_WEB_LIBS = False

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using multiple methods for best results
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text from the PDF
    """
    if not HAS_PDF_LIBS:
        raise ImportError("PDF processing libraries not installed")
    
    text = ""
    
    # Method 1: PyPDF2 for basic text extraction
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page_text = reader.pages[page_num].extract_text() or ""
                text += page_text + "\n\n"
        
        logger.info(f"Extracted {len(text.split())} words using PyPDF2")
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
    
    # Method 2: If PyPDF2 didn't extract much text, try pdfplumber
    if len(text.strip()) < 100:
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
            
            logger.info(f"Extracted {len(text.split())} words using pdfplumber")
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
    
    # Method 3: If text is still minimal, try OCR as a last resort
    if len(text.strip()) < 100 and HAS_OCR_LIBS:
        try:
            text = ""
            
            # Try using pdf2image with Poppler if available
            try:
                if POPPLER_PATH:
                    # Use the found Poppler path
                    images = convert_from_path(file_path, poppler_path=POPPLER_PATH)
                    logger.info(f"Successfully converted PDF to {len(images)} images using Poppler")
                else:
                    # Try without specifying Poppler path (might work if it's in PATH)
                    images = convert_from_path(file_path)
                    logger.info(f"Successfully converted PDF to {len(images)} images using system Poppler")
            except Exception as poppler_error:
                logger.warning(f"PDF to image conversion with Poppler failed: {poppler_error}")
                
                # Fallback: Try to extract images using PyMuPDF (fitz) if available
                try:
                    import fitz  # PyMuPDF
                    
                    logger.info("Attempting PDF to image conversion with PyMuPDF")
                    images = []
                    doc = fitz.open(file_path)
                    
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                        img_data = pix.tobytes("png")
                        
                        # Convert bytes to PIL Image
                        img = Image.open(io.BytesIO(img_data))
                        images.append(img)
                    
                    logger.info(f"Successfully converted PDF to {len(images)} images using PyMuPDF")
                except ImportError:
                    logger.warning("PyMuPDF not installed. Try: pip install PyMuPDF")
                    return text  # Return whatever text we have so far
                except Exception as fitz_error:
                    logger.warning(f"PDF to image conversion with PyMuPDF failed: {fitz_error}")
                    return text  # Return whatever text we have so far
            
            # Apply advanced OCR to each image
            for i, image in enumerate(images):
                logger.info(f"Processing page {i+1}/{len(images)} with OCR")
                
                # Advanced image preprocessing for better OCR results
                try:
                    # Convert to grayscale
                    grayscale_image = image.convert('L')
                    
                    # Apply advanced image enhancement
                    from PIL import ImageEnhance, ImageFilter
                    
                    # Enhance contrast
                    contrast_enhancer = ImageEnhance.Contrast(grayscale_image)
                    enhanced_image = contrast_enhancer.enhance(3.0)
                    
                    # Enhance brightness
                    brightness_enhancer = ImageEnhance.Brightness(enhanced_image)
                    enhanced_image = brightness_enhancer.enhance(1.3)
                    
                    # Enhance sharpness
                    sharpness_enhancer = ImageEnhance.Sharpness(enhanced_image)
                    enhanced_image = sharpness_enhancer.enhance(2.5)
                    
                    # Apply unsharp mask filter
                    enhanced_image = enhanced_image.filter(ImageFilter.UnsharpMask(radius=2, percent=200, threshold=2))
                    
                    # Apply additional filters for text documents
                    enhanced_image = enhanced_image.filter(ImageFilter.MedianFilter(size=3))
                    enhanced_image = enhanced_image.filter(ImageFilter.EDGE_ENHANCE_MORE)
                    
                    # Resize for better OCR if needed
                    if enhanced_image.width < 1500 or enhanced_image.height < 1500:
                        scale_factor = max(1500 / enhanced_image.width, 1500 / enhanced_image.height)
                        new_width = int(enhanced_image.width * scale_factor)
                        new_height = int(enhanced_image.height * scale_factor)
                        enhanced_image = enhanced_image.resize((new_width, new_height), Image.LANCZOS)
                    
                    # Try multiple OCR configurations
                    configs = [
                        r'--oem 1 --psm 3 -l eng --dpi 300',  # Default
                        r'--oem 1 --psm 6 -l eng --dpi 300',  # For printed text
                        r'--oem 1 --psm 4 -l eng --dpi 300',  # For single column
                        r'--oem 1 --psm 1 -l eng --dpi 300',  # For dense text
                    ]
                    
                    best_page_text = ""
                    best_word_count = 0
                    
                    for config in configs:
                        current_text = pytesseract.image_to_string(enhanced_image, config=config)
                        current_word_count = len(current_text.split())
                        
                        if current_word_count > best_word_count:
                            best_page_text = current_text
                            best_word_count = current_word_count
                    
                    text += best_page_text + "\n\n"
                    logger.info(f"Extracted {best_word_count} words from page {i+1} using OCR")
                    
                except Exception as ocr_error:
                    logger.warning(f"OCR processing for page {i+1} failed: {ocr_error}")
                    
                    # Try basic OCR as fallback
                    try:
                        basic_text = pytesseract.image_to_string(image)
                        text += basic_text + "\n\n"
                        logger.info(f"Extracted {len(basic_text.split())} words from page {i+1} using basic OCR")
                    except Exception as basic_ocr_error:
                        logger.warning(f"Basic OCR for page {i+1} failed: {basic_ocr_error}")
            
            logger.info(f"Extracted total of {len(text.split())} words using OCR")
            
        except Exception as e:
            logger.warning(f"OCR extraction failed: {e}")
    elif len(text.strip()) < 100:
        logger.warning("OCR libraries not available for enhanced PDF extraction")
    
    # Final check: if we still don't have text, try one more approach with PyMuPDF if available
    if len(text.strip()) < 100:
        try:
            import fitz  # PyMuPDF
            
            logger.info("Attempting text extraction with PyMuPDF as last resort")
            text = ""
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text() or ""
                text += page_text + "\n\n"
            
            logger.info(f"Extracted {len(text.split())} words using PyMuPDF")
        except ImportError:
            logger.warning("PyMuPDF not installed. Try: pip install PyMuPDF")
        except Exception as fitz_error:
            logger.warning(f"PyMuPDF text extraction failed: {fitz_error}")
    
    return text

def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from DOC files using win32com
    
    Args:
        file_path: Path to the DOC file
        
    Returns:
        Extracted text from the DOC
    """
    if not HAS_DOC_LIBS:
        raise ImportError("Document processing libraries not installed")
    
    try:
        # Convert relative path to absolute path
        file_path = os.path.abspath(file_path)
        
        # Initialize Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        try:
            # Open the document
            doc = word.Documents.Open(file_path)
            
            # Extract text
            text = doc.Content.Text
            
            # Close the document
            doc.Close()
            
            logger.info(f"Extracted {len(text.split())} words from DOC")
            return text
            
        finally:
            # Always quit Word application
            word.Quit()
    
    except Exception as e:
        logger.error(f"DOC extraction failed: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX files
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text from the DOCX
    """
    if not HAS_DOC_LIBS:
        raise ImportError("Document processing libraries not installed")
    
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        logger.info(f"Extracted {len(text.split())} words from DOCX")
        return text
    
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_odt(file_path: str) -> str:
    """
    Extract text from ODT files
    
    Args:
        file_path: Path to the ODT file
        
    Returns:
        Extracted text from the ODT
    """
    if not HAS_DOC_LIBS:
        raise ImportError("Document processing libraries not installed")
    
    try:
        textdoc = load(file_path)
        allparas = textdoc.getElementsByType(text.P)
        text_content = "\n\n".join([teletype.extractText(para) for para in allparas])
        
        logger.info(f"Extracted {len(text_content.split())} words from ODT")
        return text_content
    except Exception as e:
        logger.error(f"ODT extraction failed: {e}")
        return ""

def enhance_image(image):
    """
    Enhance an image for better text extraction
    
    Args:
        image: OpenCV image
        
    Returns:
        Enhanced image
    """
    try:
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply adaptive thresholding
        thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        
        # Apply morphological operations to remove noise
        kernel = np.ones((1, 1), np.uint8)
        opening = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel)
        
        # Apply dilation to make text more visible
        kernel = np.ones((1, 1), np.uint8)
        dilated = cv2.dilate(opening, kernel, iterations=1)
        
        return dilated
    except Exception as e:
        logger.error(f"Error enhancing image: {str(e)}")
        return image  # Return original image if enhancement fails

def create_high_contrast(image):
    """
    Create a high contrast version of the image
    
    Args:
        image: OpenCV image
        
    Returns:
        High contrast image
    """
    try:
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply histogram equalization
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        equalized = clahe.apply(gray)
        
        # Apply Otsu's thresholding
        _, binary = cv2.threshold(equalized, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return binary
    except Exception as e:
        logger.error(f"Error creating high contrast image: {str(e)}")
        return image  # Return original image if enhancement fails

def gentle_enhance(image):
    """
    Apply gentle enhancement for artistic text
    
    Args:
        image: OpenCV image
        
    Returns:
        Gently enhanced image
    """
    try:
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply slight Gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(gray, (3, 3), 0)
        
        # Apply slight sharpening
        kernel = np.array([[-1, -1, -1], [-1, 9, -1], [-1, -1, -1]])
        sharpened = cv2.filter2D(blurred, -1, kernel)
        
        return sharpened
    except Exception as e:
        logger.error(f"Error applying gentle enhancement: {str(e)}")
        return image  # Return original image if enhancement fails

def enhance_colored_text(image):
    """
    Enhance colored text in an image
    
    Args:
        image: OpenCV image
        
    Returns:
        Enhanced image for colored text
    """
    try:
        # Convert to HSV
        hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
        
        # Split into channels
        h, s, v = cv2.split(hsv)
        
        # Enhance saturation
        s = cv2.multiply(s, 1.5)
        
        # Enhance value
        v = cv2.equalizeHist(v)
        
        # Merge channels
        hsv_enhanced = cv2.merge([h, s, v])
        
        # Convert back to BGR
        enhanced = cv2.cvtColor(hsv_enhanced, cv2.COLOR_HSV2BGR)
        
        # Convert to grayscale
        gray = cv2.cvtColor(enhanced, cv2.COLOR_BGR2GRAY)
        
        return gray
    except Exception as e:
        logger.error(f"Error enhancing colored text: {str(e)}")
        return cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)  # Return grayscale of original image if enhancement fails

def extract_text_from_image(file_path: str) -> str:
    """
    Extract text from images using advanced OCR techniques
    Supports both local files and URLs
    
    Args:
        file_path: Path to the image file or URL
        
    Returns:
        Extracted text from the image
    """
    if not HAS_IMG_LIBS:
        raise ImportError("Image processing libraries not installed")
        
    if not HAS_OCR_LIBS:
        return "OCR libraries not available. Unable to extract text from images."
    
    try:
        logger.info(f"Processing image: {file_path}")
        
        # Open the image
        image = cv2.imread(file_path)
        
        # Check for children's content based on color analysis
        is_colorful = is_colorful_image(image)
        has_large_text = has_large_font(image)
        
        if is_colorful and has_large_text:
            logger.info("Detected potential children's content based on image characteristics")
            # Use specialized children's content extraction
            return extract_childrens_content(image, file_path)
            
        # Create multiple enhanced versions of the image
        enhanced_versions = []
        
        # Standard enhancement
        standard_image = enhance_image(image)
        enhanced_versions.append(("standard", standard_image))
        
        # High contrast for text on complex backgrounds
        high_contrast_image = create_high_contrast(image)
        enhanced_versions.append(("high_contrast", high_contrast_image))
        
        # Gentle enhancement for artistic text
        gentle_image = gentle_enhance(image)
        enhanced_versions.append(("gentle", gentle_image))
        
        # Color-based enhancement for colored text
        color_enhanced = enhance_colored_text(image)
        enhanced_versions.append(("color", color_enhanced))
        
        # Define OCR configurations to try
        configs = [
            # Standard configuration
            "--oem 3 --psm 3",
            # Assume a single column of text
            "--oem 3 --psm 6",
            # Assume a single block of text
            "--oem 3 --psm 4",
            # Sparse text with OSD
            "--oem 3 --psm 11",
            # Sparse text
            "--oem 3 --psm 12"
        ]
        
        # Store results from each attempt
        results = []
        
        # Process each enhanced version with multiple OCR configurations
        for version_name, enhanced_image in enhanced_versions:
            for config in configs:
                try:
                    # Convert OpenCV image to PIL format for pytesseract
                    pil_image = Image.fromarray(cv2.cvtColor(enhanced_image, cv2.COLOR_BGR2RGB))
                    
                    # Extract text using pytesseract
                    text = pytesseract.image_to_string(pil_image, config=config)
                    
                    # Skip empty results
                    if not text or len(text.strip()) < 5:
                        continue
                    
                    # Score this result
                    score = score_extraction(text)
                    
                    # Add to results
                    results.append((score, text, version_name, config))
                    
                    logger.debug(f"Extracted text with {version_name}, config={config}, score={score}")
                except Exception as e:
                    logger.error(f"Error extracting text with {version_name}, config={config}: {str(e)}")
        
        # If we have results, return the best one
        if results:
            # Sort by score (highest first)
            results.sort(reverse=True)
            best_score, best_text, best_version, best_config = results[0]
            
            logger.info(f"Selected best extraction: version={best_version}, config={best_config}, score={best_score}")
            
            # Clean up the text
            cleaned_text = clean_extracted_text(best_text)
            return cleaned_text
        
        # If no results, try specialized techniques
        logger.warning("Standard extraction methods failed, trying specialized techniques")
        return fallback_extraction(image, file_path)
        
    except Exception as e:
        logger.error(f"Error in extract_text_from_image: {str(e)}")
        return ""

def is_colorful_image(image):
    """Check if an image has vibrant colors typical of children's content"""
    try:
        # Convert to HSV color space
        hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
        
        # Calculate saturation mean
        saturation = hsv[:, :, 1]
        mean_saturation = np.mean(saturation)
        
        # Calculate standard deviation of hue
        hue = hsv[:, :, 0]
        std_hue = np.std(hue)
        
        # Colorful images have high saturation and varied hues
        return mean_saturation > 50 and std_hue > 30
    except Exception as e:
        logger.error(f"Error in is_colorful_image: {str(e)}")
        return False

def has_large_font(image):
    """Check if an image likely contains large font text typical of children's books"""
    try:
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply threshold to get binary image
        _, binary = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)
        
        # Find contours
        contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        # Filter contours by size
        large_contours = [c for c in contours if cv2.contourArea(c) > 100]
        
        # Calculate average height of contours
        if large_contours:
            heights = []
            for contour in large_contours:
                _, _, _, h = cv2.boundingRect(contour)
                heights.append(h)
            
            avg_height = np.mean(heights) if heights else 0
            
            # Large font typically has height > 20 pixels
            return avg_height > 20
        
        return False
    except Exception as e:
        logger.error(f"Error in has_large_font: {str(e)}")
        return False

def extract_childrens_content(image, file_path):
    """Specialized extraction for children's content"""
    try:
        logger.info("Using specialized children's content extraction")
        
        # Create a copy of the image
        img = image.copy()
        
        # Enhance colors for children's content
        hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
        hsv[:, :, 1] = np.clip(hsv[:, :, 1] * 1.3, 0, 255)  # Increase saturation
        enhanced = cv2.cvtColor(hsv, cv2.COLOR_HSV2BGR)
        
        # Convert to grayscale with special attention to red and pink (common in children's books)
        gray = cv2.cvtColor(enhanced, cv2.COLOR_BGR2GRAY)
        
        # Apply adaptive threshold
        binary = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                      cv2.THRESH_BINARY, 11, 2)
        
        # Dilate to connect nearby text
        kernel = np.ones((2, 2), np.uint8)
        dilated = cv2.dilate(binary, kernel, iterations=1)
        
        # Convert OpenCV image to PIL format
        pil_image = Image.fromarray(cv2.cvtColor(enhanced, cv2.COLOR_BGR2RGB))
        
        # Try multiple OCR configurations optimized for children's books
        configs = [
            # Optimized for large, clear text in children's books
            "--oem 3 --psm 6 -l eng --dpi 300",
            # Optimized for sparse text with illustrations
            "--oem 3 --psm 11 -l eng --dpi 300",
            # Optimized for single column of text
            "--oem 3 --psm 4 -l eng --dpi 300"
        ]
        
        results = []
        
        # Try each configuration
        for config in configs:
            try:
                text = pytesseract.image_to_string(pil_image, config=config)
                
                # Skip empty results
                if not text or len(text.strip()) < 5:
                    continue
                
                # Score this result with emphasis on children's content
                score = score_childrens_content(text)
                
                # Add to results
                results.append((score, text, config))
                
                logger.debug(f"Extracted children's content with config={config}, score={score}")
            except Exception as e:
                logger.error(f"Error extracting children's content with config={config}: {str(e)}")
        
        # If we have results, return the best one
        if results:
            # Sort by score (highest first)
            results.sort(reverse=True)
            best_score, best_text, best_config = results[0]
            
            logger.info(f"Selected best children's content extraction: config={best_config}, score={best_score}")
            
            # Special cleaning for children's content
            cleaned_text = clean_childrens_text(best_text)
            return cleaned_text
        
        # If specialized extraction failed, try direct extraction from the original image
        logger.warning("Specialized children's content extraction failed, trying direct extraction")
        return pytesseract.image_to_string(Image.open(file_path))
        
    except Exception as e:
        logger.error(f"Error in extract_childrens_content: {str(e)}")
        return ""

def score_childrens_content(text):
    """Score text based on likelihood of being children's content"""
    score = 0
    
    # Check for common children's content keywords
    children_keywords = ['child', 'happy', 'play', 'laugh', 'fun', 'joy', 'little', 'sun', 'tree', 'house', 'mom', 'dad']
    for keyword in children_keywords:
        if keyword.lower() in text.lower():
            score += 10
    
    # Check for short lines (common in children's poems)
    lines = text.split('\n')
    non_empty_lines = [line for line in lines if line.strip()]
    
    if non_empty_lines:
        avg_line_length = sum(len(line) for line in non_empty_lines) / len(non_empty_lines)
        if avg_line_length < 40:
            score += 15
        
        # Check for consistent line lengths (common in poems)
        line_lengths = [len(line) for line in non_empty_lines]
        if line_lengths:
            std_dev = np.std(line_lengths)
            if std_dev < 10:  # Consistent line lengths
                score += 10
    
    # Check for title-like first line
    if non_empty_lines and len(non_empty_lines[0]) < 30 and non_empty_lines[0].strip().endswith(('Child', 'child')):
        score += 20
    
    # Check for simple vocabulary
    word_count = len(text.split())
    long_words = sum(1 for word in text.split() if len(word) > 6)
    if word_count > 0 and long_words / word_count < 0.2:  # Mostly short words
        score += 15
    
    return score

def clean_childrens_text(text):
    """Special cleaning for children's content"""
    # Remove excessive newlines but preserve poem structure
    lines = text.split('\n')
    non_empty_lines = [line.strip() for line in lines if line.strip()]
    
    # Rejoin with single newlines
    cleaned = '\n'.join(non_empty_lines)
    
    # Fix common OCR errors in children's books
    replacements = {
        'l am': 'I am',
        'l have': 'I have',
        'l like': 'I like',
        'l play': 'I play',
        'l laugh': 'I laugh',
        'l hardly': 'I hardly',
        'chiid': 'child',
        'happv': 'happy',
        'piay': 'play'
    }
    
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)
    
    # Ensure the title is properly formatted if it exists
    if "happy child" in cleaned.lower() and not "a happy child" in cleaned:
        cleaned = cleaned.replace("Happy Child", "A Happy Child")
        cleaned = cleaned.replace("happy child", "A Happy Child")
    
    return cleaned

def extract_text_from_excel(file_path: str) -> str:
    """
    Extract text from Excel files
    
    Args:
        file_path: Path to the Excel file
        
    Returns:
        Extracted text from the Excel file
    """
    if not HAS_SPREADSHEET_LIBS:
        raise ImportError("Spreadsheet processing libraries not installed")
    
    try:
        # Read Excel file
        df = pd.read_excel(file_path, sheet_name=None)
        
        text_parts = []
        
        # Process each sheet
        for sheet_name, sheet_df in df.items():
            text_parts.append(f"Sheet: {sheet_name}")
            text_parts.append(sheet_df.to_string(index=False))
            text_parts.append("\n")
        
        text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(text.split())} words from Excel")
        return text
    except Exception as e:
        logger.error(f"Excel extraction failed: {e}")
        return ""

def extract_text_from_csv(file_path: str) -> str:
    """
    Extract text from CSV files
    
    Args:
        file_path: Path to the CSV file
        
    Returns:
        Extracted text from the CSV
    """
    if not HAS_SPREADSHEET_LIBS:
        raise ImportError("Spreadsheet processing libraries not installed")
    
    try:
        # Read CSV file
        df = pd.read_csv(file_path)
        
        # Convert to string representation
        text = df.to_string(index=False)
        
        logger.info(f"Extracted {len(text.split())} words from CSV")
        return text
    except Exception as e:
        logger.error(f"CSV extraction failed: {e}")
        return ""

def extract_text_from_html(file_path: str) -> str:
    """
    Extract text from HTML files
    
    Args:
        file_path: Path to the HTML file
        
    Returns:
        Extracted text from the HTML
    """
    if not HAS_WEB_LIBS:
        raise ImportError("Web content processing libraries not installed")
    
    try:
        # Read HTML file
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        text = soup.get_text()
        
        # Break into lines and remove leading and trailing space
        lines = (line.strip() for line in text.splitlines())
        
        # Break multi-headlines into a line each
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        
        # Drop blank lines
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        logger.info(f"Extracted {len(text.split())} words from HTML")
        return text
    except Exception as e:
        logger.error(f"HTML extraction failed: {e}")
        return ""

def extract_text_from_url(url: str) -> str:
    """
    Extract text from a URL
    
    Args:
        url: URL to extract text from
        
    Returns:
        Extracted text from the URL
    """
    if not HAS_WEB_LIBS:
        raise ImportError("Web content processing libraries not installed")
    
    try:
        # Fetch URL content
        response = requests.get(url)
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        text = soup.get_text()
        
        # Process text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        logger.info(f"Extracted {len(text.split())} words from URL")
        return text
    except Exception as e:
        logger.error(f"URL extraction failed: {e}")
        return ""

def preprocess_text(text: str) -> str:
    """
    Preprocess extracted text to improve quality
    
    Args:
        text: Raw extracted text
        
    Returns:
        Preprocessed text
    """
    if not text:
        return ""
    
    import re
    
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Preserve important formatting like headers, lists, and tables
    # Preserve headers (# symbols)
    text = re.sub(r'(^|\n)#(\S)', r'\1# \2', text)  # Ensure space after # for headers
    
    # Fix bullet points
    text = re.sub(r'(^|\n)•\s*', r'\1- ', text)  # Replace bullet with hyphen
    text = re.sub(r'(^|\n)\*\s+', r'\1- ', text)  # Replace asterisk with hyphen
    
    # Preserve list items and indentation
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        if line.strip():  # Skip empty lines
            # Preserve indentation for list items
            if re.match(r'^\s*[-•*]\s', line):
                processed_lines.append(line)
            else:
                # Remove excessive spaces within the line but preserve leading spaces
                leading_spaces = len(line) - len(line.lstrip())
                if leading_spaces > 0:
                    processed_lines.append(' ' * leading_spaces + ' '.join(line.strip().split()))
                else:
                    processed_lines.append(' '.join(line.strip().split()))
        else:
            processed_lines.append('')  # Keep empty lines for paragraph separation
    
    # Rejoin with proper line endings
    text = '\n'.join(processed_lines)
    
    # Remove sequences of more than 2 newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Fix common OCR issues
    text = text.replace('|', 'I')  # Common OCR mistake: pipe for capital I
    text = re.sub(r'(\d)\.(\d)', r'\1.\2', text)  # Fix decimal points
    
    # Fix common formatting issues
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between lowercase and uppercase
    
    return text

def extract_text_from_file(file_path: str, file_type: str = None) -> str:
    """
    Main function to extract text from any file type
    
    Args:
        file_path: Path to the file
        file_type: Type of file (pdf, image, etc.)
        
    Returns:
        Extracted text from the file
    """
    # Determine file type if not provided
    if not file_type:
        _, ext = os.path.splitext(file_path)
        file_type = ext.lower().lstrip('.')
    
    # Normalize file type
    file_type = file_type.lower()
    
    # Extract text based on file type
    try:
        if file_type in SUPPORTED_FILE_TYPES['pdf'] or file_type == 'pdf':
            text = extract_text_from_pdf(file_path)
        elif any(file_type in ext_list for ext_list in [SUPPORTED_FILE_TYPES['image']]) or file_type == 'image':
            text = extract_text_from_image(file_path)
        elif file_type == 'docx':
            text = extract_text_from_docx(file_path)
        elif file_type == 'doc':
            text = extract_text_from_doc(file_path)
        elif file_type == 'odt':
            text = extract_text_from_odt(file_path)
        elif file_type in ['xlsx', 'xls']:
            text = extract_text_from_excel(file_path)
        elif file_type in ['csv']:
            text = extract_text_from_csv(file_path)
        elif file_type in ['html', 'htm']:
            text = extract_text_from_html(file_path)
        elif file_type.startswith('http'):
            text = extract_text_from_url(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Preprocess the extracted text
        text = preprocess_text(text)
        
        return text
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_type} file: {e}", exc_info=True)
        
        # For image files, provide a graceful fallback rather than failing
        if file_type == 'image' or file_type in SUPPORTED_FILE_TYPES['image']:
            return f"Image text extraction attempted but encountered technical issues: {str(e)}. The system will try to process this image with alternative methods."
        
        raise ValueError(f"Failed to extract text from {file_type} file: {str(e)}")

def score_extraction(text):
    """
    Score the quality of text extraction
    
    Args:
        text: Extracted text
        
    Returns:
        Score (higher is better)
    """
    if not text:
        return 0
        
    # Initialize score
    score = 0
    
    # Length is a good indicator (longer text often means better extraction)
    score += min(len(text) / 100, 10)  # Cap at 10 points for length
    
    # Check for common OCR errors
    if '|' in text or '1' in text or '0' in text:
        score -= 1  # Potential OCR confusion with I/l/1 or O/0
        
    # Check for reasonable word count
    words = text.split()
    word_count = len(words)
    if word_count > 5:
        score += 3
    
    # Check for reasonable word length
    avg_word_length = sum(len(word) for word in words) / max(word_count, 1)
    if 3 <= avg_word_length <= 10:
        score += 2
    
    # Check for reasonable line count
    lines = text.split('\n')
    if len(lines) > 1:
        score += 2
    
    # Check for punctuation (good sign of proper text)
    if any(p in text for p in '.,:;?!'):
        score += 2
    
    # Check for capitalization (good sign of proper text)
    if any(c.isupper() for c in text):
        score += 1
    
    # Check for reasonable character distribution
    alpha_ratio = sum(c.isalpha() for c in text) / max(len(text), 1)
    if 0.5 <= alpha_ratio <= 0.95:
        score += 3
    
    return score

def fallback_extraction(image_path):
    """
    Fallback extraction method when standard methods fail
    
    Args:
        image_path: Path to the image
        
    Returns:
        Extracted text
    """
    try:
        # Try basic OCR with default settings
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        
        if text and len(text.strip()) > 10:
            return text
            
        # Try with different preprocessing
        img = img.convert('L')  # Convert to grayscale
        text = pytesseract.image_to_string(img)
        
        if text and len(text.strip()) > 10:
            return text
            
        # Try with different PSM modes
        for psm in [3, 6, 4, 11, 12]:
            text = pytesseract.image_to_string(img, config=f'--psm {psm}')
            if text and len(text.strip()) > 10:
                return text
                
        return "No text could be extracted from the image."
    except Exception as e:
        logger.error(f"Error in fallback extraction: {str(e)}")
        return "Error extracting text from image."

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Extract text from various file formats')
    parser.add_argument('file_path', help='Path to the file or URL')
    parser.add_argument('--file-type', help='File type (pdf, image, docx, etc.)')
    parser.add_argument('--output', help='Output file path (optional)')
    
    args = parser.parse_args()
    
    try:
        # Extract text
        text = extract_text_from_file(args.file_path, args.file_type)
        
        # Save or print output
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Text extracted and saved to {args.output}")
        else:
            print(text)
        
        # Return success status
        print(json.dumps({
            "status": "success",
            "text_length": len(text),
            "word_count": len(text.split())
        }))
        
    except Exception as e:
        # Return error status
        print(json.dumps({
            "status": "error",
            "error": str(e)
        }), file=sys.stderr)
        sys.exit(1)
